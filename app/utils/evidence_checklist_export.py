"""
Evidence checklist export — PDF and DOCX formats.

Generates a formatted evidence request document that can be sent to the client
before the assessment begins.
"""

from datetime import datetime, timezone
from io import BytesIO

from fpdf import FPDF

from app.utils.pdf_export import (
    NAVY, DARK_TEXT, MID_TEXT, LIGHT_TEXT, WHITE, CARD_BG, DIVIDER,
    PW, PM, CW, S,
)


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------

def generate_evidence_checklist_pdf(
    company_name: str,
    checklist: list[dict],
    flags: dict,
) -> bytes:
    """Generate a professional evidence request PDF."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.alias_nb_pages()

    date_str = datetime.now(timezone.utc).strftime("%d %B %Y")

    # --- Cover ---
    pdf.add_page()
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, PW, 55, style="F")

    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*WHITE)
    pdf.set_y(15)
    pdf.cell(0, 10, text=S("Evidence Request"), align="C")
    pdf.ln(8)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 7, text=S("DPDPA Compliance Gap Assessment"), align="C")
    pdf.ln(7)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, text=S(f"Prepared for: {company_name}  |  {date_str}"), align="C")

    pdf.set_y(65)
    pdf.set_text_color(*DARK_TEXT)

    # --- Introduction ---
    pdf.set_font("Helvetica", "", 10)
    pdf.set_fill_color(*CARD_BG)
    pdf.rect(PM, pdf.get_y(), CW, 22, style="F")
    pdf.set_x(PM + 5)
    pdf.set_y(pdf.get_y() + 4)
    pdf.multi_cell(
        CW - 10, 6,
        text=S(
            "Please provide the documents listed below to enable a thorough DPDPA gap assessment. "
            "Required documents are essential for the assessment. Recommended documents improve "
            "coverage and may reduce the number of follow-up questions."
        ),
    )
    pdf.ln(6)

    # --- Scope flags ---
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*NAVY)
    pdf.cell(0, 6, text=S("ASSESSMENT SCOPE"), ln=True)
    pdf.set_draw_color(*DIVIDER)
    pdf.line(PM, pdf.get_y(), PM + CW, pdf.get_y())
    pdf.ln(3)

    flag_labels = [
        ("Cross-border transfers", flags.get("cross_border", False)),
        ("Children's data", flags.get("children", False)),
        ("SDF obligations", flags.get("sdf", False)),
        ("Third-party processors", flags.get("processors", False)),
    ]

    pdf.set_font("Helvetica", "", 9)
    col_w = CW / 2
    items_per_row = 2
    for i, (label, active) in enumerate(flag_labels):
        if i % items_per_row == 0 and i > 0:
            pdf.ln(7)
        x_offset = PM + (i % items_per_row) * col_w
        pdf.set_x(x_offset)
        color = (39, 174, 96) if active else (189, 195, 199)
        pdf.set_text_color(*color)
        prefix = S("[ACTIVE]  ") if active else S("[N/A]  ")
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(20, 6, text=prefix)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(col_w - 25, 6, text=S(label))
    pdf.ln(12)

    # --- Checklist ---
    required = [c for c in checklist if c["required"]]
    recommended = [c for c in checklist if not c["required"]]

    def _section(title: str, items: list[dict], required: bool):
        if not items:
            return
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*NAVY)
        pdf.cell(0, 6, text=S(title.upper()), ln=True)
        pdf.set_draw_color(*DIVIDER)
        pdf.line(PM, pdf.get_y(), PM + CW, pdf.get_y())
        pdf.ln(4)

        for item in items:
            # Check if we need a new page
            if pdf.get_y() > 260:
                pdf.add_page()
                pdf.ln(5)

            # Checkbox
            box_y = pdf.get_y()
            pdf.set_draw_color(*NAVY if required else (180, 180, 180))
            pdf.set_line_width(0.4)
            if required:
                pdf.rect(PM, box_y, 5, 5)
            else:
                # Dashed look for recommended — just draw a lighter box
                pdf.set_draw_color(200, 200, 200)
                pdf.rect(PM, box_y, 5, 5)

            # Label
            pdf.set_x(PM + 8)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*DARK_TEXT)
            pdf.cell(CW - 8, 5, text=S(item["label"]))
            pdf.ln(5)

            # Reason
            pdf.set_x(PM + 8)
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(*MID_TEXT)
            pdf.multi_cell(CW - 8, 5, text=S(item["reason"]))

            # Requirement tags
            if item.get("maps_to"):
                pdf.set_x(PM + 8)
                tags = "  ".join(item["maps_to"])
                pdf.set_font("Helvetica", "", 8)
                pdf.set_text_color(*LIGHT_TEXT)
                pdf.cell(CW - 8, 4, text=S(tags))
                pdf.ln(4)

            pdf.ln(3)

        pdf.ln(4)

    _section("Required Documents", required, required=True)
    _section("Recommended Documents", recommended, required=False)

    # --- Footer note ---
    if pdf.get_y() > 250:
        pdf.add_page()

    pdf.set_y(pdf.get_y() + 6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(*LIGHT_TEXT)
    pdf.multi_cell(CW, 5, text=S(
        "This evidence request was generated by CyberAssess. "
        "Documents should be provided in PDF or DOCX format where possible. "
        "If a document does not exist, please note that explicitly — gaps are part of the assessment."
    ), align="C")

    return bytes(pdf.output())


# ---------------------------------------------------------------------------
# DOCX Export
# ---------------------------------------------------------------------------

def generate_evidence_checklist_docx(
    company_name: str,
    checklist: list[dict],
    flags: dict,
) -> bytes:
    """Generate an editable DOCX evidence request document."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    date_str = datetime.now(timezone.utc).strftime("%d %B %Y")

    # Title
    title = doc.add_heading("Evidence Request", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"DPDPA Compliance Gap Assessment\nPrepared for: {company_name}  |  {date_str}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Introduction
    intro = doc.add_paragraph(
        "Please provide the documents listed below to enable a thorough DPDPA gap assessment. "
        "Required documents are essential for the assessment. Recommended documents improve "
        "coverage and may reduce the number of follow-up questions. "
        "If a document does not exist, please note that explicitly — gaps are part of the assessment."
    )
    intro.style = "Normal"

    doc.add_paragraph()

    # Scope flags
    scope_heading = doc.add_heading("Assessment Scope", level=2)
    flag_labels = [
        ("Cross-border transfers", flags.get("cross_border", False)),
        ("Children's data", flags.get("children", False)),
        ("SDF obligations", flags.get("sdf", False)),
        ("Third-party processors", flags.get("processors", False)),
    ]
    for label, active in flag_labels:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{'[ACTIVE]' if active else '[Not applicable]'}  {label}")
        run.bold = active

    doc.add_paragraph()

    # Checklist sections
    required = [c for c in checklist if c["required"]]
    recommended = [c for c in checklist if not c["required"]]

    def _section(title: str, items: list[dict]):
        if not items:
            return
        doc.add_heading(title, level=2)
        for i, item in enumerate(items, 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(item["label"]).bold = True
            doc.add_paragraph(item["reason"]).style = "Normal"
            tags = "  ".join(item["maps_to"])
            req_p = doc.add_paragraph(f"Requirements: {tags}")
            req_p.runs[0].italic = True
            doc.add_paragraph()  # spacer

    _section("Required Documents", required)
    _section("Recommended Documents", recommended)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
