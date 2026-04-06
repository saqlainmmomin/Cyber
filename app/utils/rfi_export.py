"""
RFI (Request for Information) export — PDF and DOCX formats.

PDF: Professional letterhead style using fpdf2.
DOCX: Editable format using python-docx.
"""

import json
from datetime import datetime, timezone
from io import BytesIO

from fpdf import FPDF

# Reuse palette and helpers from main PDF export
from app.utils.pdf_export import (
    NAVY, DARK_TEXT, MID_TEXT, LIGHT_TEXT, WHITE, CARD_BG, DIVIDER,
    PW, PM, CW, S,
)

# Priority colors for RFI
RFI_PRIORITY_COLORS = {
    "Critical": (231, 76, 60),
    "High": (230, 126, 34),
    "Medium": (241, 196, 15),
    "Low": (39, 174, 96),
}


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------

def generate_rfi_pdf(
    title: str,
    company_name: str,
    introduction: str,
    evidence_items: list[dict],
    response_instructions: str,
    generated_at: datetime | None = None,
) -> bytes:
    """Generate a professional RFI PDF document."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.alias_nb_pages()

    date_str = (generated_at or datetime.now(timezone.utc)).strftime("%d %B %Y")

    # --- Cover Page ---
    pdf.add_page()
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, PW, 60, style="F")

    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(*WHITE)
    pdf.set_y(18)
    pdf.cell(0, 10, text=S("Request for Information"), align="C")
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, text=S("DPDPA Compliance Gap Assessment"), align="C")

    pdf.set_y(70)
    pdf.set_text_color(*DARK_TEXT)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, text=S(company_name), align="C")
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*MID_TEXT)
    pdf.cell(0, 6, text=S(f"Date: {date_str}"), align="C")
    pdf.ln(6)
    pdf.cell(0, 6, text=S(f"Reference: RFI-{company_name[:20].upper().replace(' ', '-')}"), align="C")

    # Summary box
    pdf.ln(15)
    critical = sum(1 for i in evidence_items if i.get("priority") == "Critical")
    high = sum(1 for i in evidence_items if i.get("priority") == "High")
    medium = sum(1 for i in evidence_items if i.get("priority") == "Medium")

    pdf.set_fill_color(*CARD_BG)
    pdf.rect(PM, pdf.get_y(), CW, 30, style="F")
    y = pdf.get_y() + 5
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*NAVY)
    pdf.set_xy(PM + 5, y)
    pdf.cell(0, 6, text=S(f"Total Evidence Items: {len(evidence_items)}"))
    pdf.set_xy(PM + 5, y + 8)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*MID_TEXT)

    summary_parts = []
    if critical:
        summary_parts.append(f"{critical} Critical")
    if high:
        summary_parts.append(f"{high} High")
    if medium:
        summary_parts.append(f"{medium} Medium")
    pdf.cell(0, 6, text=S(" | ".join(summary_parts) if summary_parts else "No priority items"))
    pdf.set_y(pdf.get_y() + 25)

    # Confidentiality notice
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(*LIGHT_TEXT)
    pdf.multi_cell(CW, 4, text=S(
        "CONFIDENTIAL: This document contains findings from a DPDPA compliance assessment. "
        "Distribution should be limited to authorized personnel."
    ), align="C")

    _rfi_footer(pdf, company_name)

    # --- Introduction Page ---
    pdf.add_page()
    _rfi_header(pdf, "Introduction")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.multi_cell(CW, 5, text=S(introduction))
    _rfi_footer(pdf, company_name)

    # --- Evidence Items Pages ---
    # Group by chapter
    chapters = {}
    for item in evidence_items:
        ch = item.get("chapter", "General")
        chapters.setdefault(ch, []).append(item)

    for chapter_name, items in chapters.items():
        pdf.add_page()
        _rfi_header(pdf, f"Evidence Items - {chapter_name}")

        for item in items:
            _render_evidence_item(pdf, item, company_name)

        _rfi_footer(pdf, company_name)

    # --- Response Instructions Page ---
    pdf.add_page()
    _rfi_header(pdf, "Response Instructions")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.multi_cell(CW, 5, text=S(response_instructions))
    _rfi_footer(pdf, company_name)

    return pdf.output()


def _render_evidence_item(pdf: FPDF, item: dict, company_name: str):
    """Render a single evidence request item as a card."""
    if pdf.get_y() > 240:
        pdf.add_page()
        _rfi_header(pdf, "Evidence Items (continued)")
        _rfi_footer(pdf, company_name)

    y_start = pdf.get_y()

    # Item header bar
    priority = item.get("priority", "Medium")
    color = RFI_PRIORITY_COLORS.get(priority, (150, 150, 150))
    pdf.set_fill_color(*color)
    pdf.rect(PM, y_start, 3, 40, style="F")  # Left color bar

    pdf.set_fill_color(*CARD_BG)
    pdf.rect(PM + 3, y_start, CW - 3, 40, style="F")

    # Item ID + Priority
    pdf.set_xy(PM + 6, y_start + 2)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*NAVY)
    pdf.cell(30, 5, text=S(item.get("item_id", "")))

    pdf.set_font("Helvetica", "B", 8)
    pdf.set_text_color(*color)
    pdf.cell(20, 5, text=S(priority))

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(*LIGHT_TEXT)
    pdf.cell(0, 5, text=S(f"{item.get('requirement_id', '')} | {item.get('dpdpa_section', '')}"))

    # Requirement title
    pdf.set_xy(PM + 6, y_start + 8)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.multi_cell(CW - 12, 4, text=S(item.get("requirement_title", "")[:100]))

    # Current status
    pdf.set_x(PM + 6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(*MID_TEXT)
    status_text = item.get("current_status", "")[:150]
    pdf.multi_cell(CW - 12, 4, text=S(f"Status: {status_text}"))

    # Evidence requested
    pdf.set_x(PM + 6)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(*DARK_TEXT)
    evidence_text = item.get("evidence_requested", "")[:300]
    pdf.multi_cell(CW - 12, 4, text=S(f"Evidence Requested: {evidence_text}"))

    # Deadline
    pdf.set_x(PM + 6)
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(*LIGHT_TEXT)
    deadline = item.get("deadline_weeks", 3)
    pdf.cell(0, 4, text=S(f"Suggested deadline: {deadline} weeks"))

    pdf.set_y(pdf.get_y() + 8)


def _rfi_header(pdf: FPDF, section_title: str):
    """Draw RFI header bar."""
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, PW, 12, style="F")
    pdf.set_font("Helvetica", "B", 7)
    pdf.set_text_color(*WHITE)
    pdf.text(PM, 8, S(f"RFI | {section_title}"))
    pdf.set_y(16)


def _rfi_footer(pdf: FPDF, company_name: str):
    """Draw RFI footer."""
    pdf.set_y(-15)
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(*LIGHT_TEXT)
    pdf.cell(0, 5, text=S(f"CONFIDENTIAL  |  RFI  |  {company_name}"), align="L")
    pdf.cell(0, 5, text=f"Page {pdf.page_no()}/{{nb}}", align="R")


# ---------------------------------------------------------------------------
# DOCX Export
# ---------------------------------------------------------------------------

def generate_rfi_docx(
    title: str,
    company_name: str,
    introduction: str,
    evidence_items: list[dict],
    response_instructions: str,
    generated_at: datetime | None = None,
) -> bytes:
    """Generate a professional RFI DOCX document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    date_str = (generated_at or datetime.now(timezone.utc)).strftime("%d %B %Y")

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title_p = doc.add_heading(title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(f"Date: {date_str}\nOrganization: {company_name}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    # Confidentiality
    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_p.add_run("CONFIDENTIAL — Distribution limited to authorized personnel")
    conf_run.font.size = Pt(8)
    conf_run.font.italic = True
    conf_run.font.color.rgb = RGBColor(150, 150, 150)

    # Summary
    critical = sum(1 for i in evidence_items if i.get("priority") == "Critical")
    high = sum(1 for i in evidence_items if i.get("priority") == "High")
    summary_p = doc.add_paragraph()
    summary_p.add_run(f"Total Evidence Items: {len(evidence_items)}").bold = True
    if critical:
        summary_p.add_run(f"  |  {critical} Critical").font.color.rgb = RGBColor(231, 76, 60)
    if high:
        summary_p.add_run(f"  |  {high} High").font.color.rgb = RGBColor(230, 126, 34)

    doc.add_page_break()

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(introduction)

    doc.add_page_break()

    # Evidence Items Table
    doc.add_heading("Evidence Items", level=1)

    # Group by chapter
    chapters = {}
    for item in evidence_items:
        ch = item.get("chapter", "General")
        chapters.setdefault(ch, []).append(item)

    for chapter_name, items in chapters.items():
        doc.add_heading(chapter_name, level=2)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ["Item ID", "Requirement", "Priority", "Current Status", "Evidence Requested"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)

        # Data rows
        for item in items:
            row = table.add_row()
            row.cells[0].text = item.get("item_id", "")
            row.cells[1].text = f"{item.get('requirement_title', '')}\n({item.get('requirement_id', '')})"
            row.cells[2].text = item.get("priority", "Medium")
            row.cells[3].text = item.get("current_status", "")[:200]
            row.cells[4].text = item.get("evidence_requested", "")[:300]

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        doc.add_paragraph()  # spacer

    doc.add_page_break()

    # Response Instructions
    doc.add_heading("Response Instructions", level=1)
    doc.add_paragraph(response_instructions)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
