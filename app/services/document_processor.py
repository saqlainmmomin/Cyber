"""
Document text extraction for PDF, DOCX, and image files.

Supported:
- PDF: text extraction via pdfplumber
- DOCX: text extraction via python-docx
- PNG / JPG / JPEG / WEBP: content description via Claude vision API
"""

import base64
import os
import re
import uuid

import anthropic
import pdfplumber
from docx import Document

from app.config import settings

# Media type map for Claude vision
_IMAGE_MEDIA_TYPES = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
}


def save_upload(assessment_id: str, filename: str, content: bytes) -> str:
    """Save uploaded file to disk and return the file path."""
    dest_dir = os.path.join(settings.upload_dir, assessment_id)
    os.makedirs(dest_dir, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex[:8]}_{filename}"
    file_path = os.path.join(dest_dir, safe_name)
    with open(file_path, "wb") as f:
        f.write(content)
    return file_path


def extract_text(file_path: str, file_type: str) -> str:
    """Extract text from a PDF, DOCX, or image file."""
    if file_type == "pdf":
        return _extract_pdf(file_path)
    elif file_type == "docx":
        return _extract_docx(file_path)
    elif file_type in _IMAGE_MEDIA_TYPES:
        return _extract_image(file_path, file_type)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber (handles tables well)."""
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
            # Also extract tables as text
            for table in page.extract_tables():
                rows = []
                for row in table:
                    cells = [str(c) if c else "" for c in row]
                    rows.append(" | ".join(cells))
                if rows:
                    pages.append("\n".join(rows))
    full_text = "\n\n".join(pages)
    return _truncate(full_text)


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX including paragraphs and tables."""
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    full_text = "\n".join(parts)
    return _truncate(full_text)


def _truncate(text: str) -> str:
    """Truncate text to max_document_words."""
    # Clean up whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    words = text.split()
    if len(words) > settings.max_document_words:
        words = words[: settings.max_document_words]
        return " ".join(words) + "\n\n[... truncated to first {} words ...]".format(
            settings.max_document_words
        )
    return text


def extract_relevant_sections(text: str, focus_areas: list[str], max_words: int = 5000) -> str:
    """
    Extract sections of a document most relevant to the given focus areas.

    Uses section headers as anchors instead of cutting mid-sentence.
    Falls back to word-count truncation if no headers are detected.
    """
    # Split by common section headers (lines that look like headings)
    header_pattern = re.compile(
        r"^(?:\d+[\.\)]\s*|#{1,3}\s*|[A-Z][A-Z\s]{3,}:?\s*$)",
        re.MULTILINE,
    )
    headers = list(header_pattern.finditer(text))

    if len(headers) < 2:
        # No detectable sections — fall back to word-count truncation
        return _truncate_to_words(text, max_words)

    # Build sections from header positions
    sections = []
    for i, match in enumerate(headers):
        start = match.start()
        end = headers[i + 1].start() if i + 1 < len(headers) else len(text)
        section_text = text[start:end].strip()
        sections.append(section_text)

    # Score each section by keyword overlap with focus areas
    focus_lower = [f.lower() for f in focus_areas]
    scored = []
    for section in sections:
        section_lower = section.lower()
        score = sum(1 for kw in focus_lower if kw in section_lower)
        scored.append((score, section))

    # Sort by relevance (highest score first), then assemble up to max_words
    scored.sort(key=lambda x: x[0], reverse=True)

    result_parts = []
    total_words = 0
    for score, section in scored:
        words = section.split()
        if total_words + len(words) > max_words:
            remaining = max_words - total_words
            if remaining > 50:  # Only include if we can fit a meaningful chunk
                result_parts.append(" ".join(words[:remaining]))
                total_words += remaining
            break
        result_parts.append(section)
        total_words += len(words)

    return "\n\n".join(result_parts)


def _truncate_to_words(text: str, max_words: int) -> str:
    """Truncate text to max_words, trying to end at a sentence boundary."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    words = text.split()
    if len(words) <= max_words:
        return text
    truncated = " ".join(words[:max_words])
    # Try to end at last sentence boundary
    last_period = truncated.rfind(".")
    if last_period > len(truncated) * 0.8:
        truncated = truncated[: last_period + 1]
    return truncated + "\n\n[... truncated ...]"


def _extract_image(file_path: str, file_type: str) -> str:
    """
    Extract compliance-relevant content from an image using Claude vision.

    Returns a structured text description of what is visible in the screenshot
    so it can be fed into the evidence extraction pipeline.
    """
    media_type = _IMAGE_MEDIA_TYPES[file_type]
    with open(file_path, "rb") as f:
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")

    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
    message = client.messages.create(
        model=settings.claude_model,
        max_tokens=1500,
        system=(
            "You are a compliance document analyst. When shown a screenshot or image, "
            "extract and transcribe all visible text exactly as it appears. Then add a "
            "brief structured summary of what the image shows in the context of data "
            "protection compliance (e.g. consent screen, privacy policy excerpt, cookie "
            "banner, breach log, data flow diagram). Format: first transcribe the text "
            "verbatim under 'VISIBLE TEXT:', then add 'SUMMARY:' with 2-3 sentences "
            "describing what compliance-relevant controls or gaps the image reveals."
        ),
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            "Please transcribe all visible text from this image and provide "
                            "a compliance-focused summary of what it shows."
                        ),
                    },
                ],
            }
        ],
    )

    description = message.content[0].text
    filename = os.path.basename(file_path)
    return f"[Screenshot: {filename}]\n\n{description}"


def detect_file_type(filename: str) -> str | None:
    """Detect file type from extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return "pdf"
    elif ext == "docx":
        return "docx"
    elif ext in _IMAGE_MEDIA_TYPES:
        return ext
    return None
